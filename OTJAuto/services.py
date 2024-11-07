import os
from models import Activity
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

# You can add your own modules and respective KSBs here
MODULE_KSBS = {
    "Business Organisation and Strategy": ["C4", "C5", "C8", "C12", "C14", "C15", "C18", "C19", "C26"],
    "Management of Digital Resources": ["C18", "C26", "C9", "C15"],
    "Project Management and Project Proposal": ["C18", "C19"],
    "Agile Web Applications": ["SE1", "SE2", "SE3", "SE4", "SE5", "SE6", "SE7", "SE8", "SE9", "SE10", "SE11", "SE12", "SE13", "SE14"],
    "UML and Design Patterns": ["C2", "C16", "C18", "C24"],
    "Negotiated Work Based Project": ["C6", "C9", "C10", "C11", "C16", "C17", "C18", "C19", "C20", "C21", "C22", "C23", "C24", "C25", "C26", "C27", "C28", "C29", "C30", "C31"]
}

class ActivityService:
    def __init__(self):
        # You can add your own student ID and name here
        self.activities = []
        self.student_id = "A00000Z"
        self.first_name = "Change"
        self.surname = "Me"
    
    def add_activity(self, date, start_time, duration, description, module_names):
        ksbs = []
        for module_name in module_names:
            ksbs.extend(MODULE_KSBS.get(module_name, []))

        ksbs = list(set(ksbs))

        activity = Activity(date, start_time, duration, description, ', '.join(module_names), ksbs)
        self.activities.append(activity)
        print(f"Added activity: {activity}")

    def list_activities(self):
        for idx, activity in enumerate(self.activities, start=1):
            print(f"{idx}. {activity}")

    def generate_report(self, report_path='reports/activity_report.docx'):
        doc = Document()
    
        title = doc.add_paragraph("Apprenticeship Log")
        title.runs[0].font.size = Pt(24)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        info_table = doc.add_table(rows=3, cols=2)
        info_table.style = 'Table Grid'
        info_table.cell(0, 0).text = "Student ID"
        info_table.cell(0, 1).text = self.student_id
        info_table.cell(1, 0).text = "Apprentice First Name"
        info_table.cell(1, 1).text = self.first_name
        info_table.cell(2, 0).text = "Apprentice Surname"
        info_table.cell(2, 1).text = self.surname

        doc.add_paragraph("\n") 

        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        header_cells = table.rows[0].cells
        header_cells[0].text = "Date/Time"
        header_cells[1].text = "Activity/Module"
        header_cells[2].text = "Reflection/What did you learn?"
        header_cells[3].text = "Hours"
        header_cells[4].text = "KSBs Covered"

        total_hours = 0

        ksb_count = {ksb: 0 for module in MODULE_KSBS.values() for ksb in module}

        for entry in self.activities:
            row_cells = table.add_row().cells
            row_cells[0].text = f"{entry.date.strftime('%Y-%m-%d')} {entry.start_time.strftime('%H:%M')} - {entry.end_time}"
            row_cells[1].text = entry.module_names
            row_cells[2].text = entry.description
            row_cells[3].text = str(entry.duration)
            row_cells[4].text = ", ".join(entry.ksbs)

            total_hours += entry.duration

            for ksb in entry.ksbs:
                if ksb in ksb_count:
                    ksb_count[ksb] += 1

            doc.add_paragraph("\n")

        total_row_cells = table.add_row().cells
        total_row_cells[0].merge(total_row_cells[3])
        total_row_cells[0].text = "Total hours"
        total_row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        total_row_cells[4].text = str(total_hours)

        doc.add_paragraph("\nKSB Statistics:")

        ksb_table = doc.add_table(rows=1, cols=3)
        ksb_table.style = 'Table Grid'
        
        ksb_table.cell(0, 0).text = "KSB"
        ksb_table.cell(0, 1).text = "Frequency"
        ksb_table.cell(0, 2).text = "Highlight"

        sorted_ksb_count = sorted(ksb_count.items(), key=lambda item: item[1], reverse=True)

        for ksb, count in sorted_ksb_count:
            if count > 0:
                row = ksb_table.add_row().cells
                row[0].text = ksb
                row[1].text = str(count)
                row[2].text = "Most Worked-on" if count == sorted_ksb_count[0][1] else ""

        os.makedirs(os.path.dirname(report_path), exist_ok=True)
        doc.save(report_path)
        print(f"Report generated at {report_path}")
